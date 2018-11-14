#! /usr/bin/env/ python3

'''This is a test module, to present the results from the plagiarism.py
program. It first sets up a dummy doc object, as if it had been produced by
plagiarism.py. It should then save the output as a FORMATTED Word document,
with the plagiarised text highlighted'''
import docx
import os
import get_texts
import shutil


def fake_results():
    original_essay = 'this is the good text. This, however, is the plagerised
    text - to be highlighted. This is good text.'
    shortlist = ['This, however, is', 'plagerised text - to be highlighted.']
    return(original_essay, shortlist)


def results(shortlist, original_essay):
    old_index = 0
    out_document = docx.Document()
    for result in shortlist:
        index = original_essay.index(result)
        para = out_document.add_paragraph()
        para.add_run(original_essay[old_index: index])
    out_document.save('out_document.docx')


def main():
    original_essay, shortlist = fake_results()
    results(shortlist, original_essay)
    os.system('libreoffice out_document.docx')
    os.system('rm out_document.docx')


if __name__ == '__main__':
    main()


'''notes on docx module from 'automate the boring stuff'''
#    import docx
#    doc = docx.Document()
#    doc.add_paragraph('Hello world!')
#    paraObj1 = doc.add_paragraph('This is a second paragraph.')
#    paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
#    paraObj1.add_run(' This text is being added to the second paragraph.')
#    doc.save('multipleParagraphs.docx')
