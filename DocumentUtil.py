import docx
import tkinter as tk
from tkinter import filedialog


class DocumentReader:
    """
    Provides tools to open docx files and make
    the text available in different formats
    """
    def __init__(self, title=''):
        self.title = title
        self.doc_object = None
        self.paragraphs = []
        self.whole_text = None

    def file_picker():
        """
        A gui file picker. If a title is given
        the title will be visible to aid selection
        """
        # add title functionality
        root = tk.Tk()
        root.withdraw()
        filePath = filedialog.askopenfilename()
        return filePath

    def get_as_paragraphs(self):
        """
        Gets all paragraphs in the file
        text and retuns them in a list
        """
        Print(f'select {self.title} text')
        filePath = self.file_picker
        self.doc_object = docx.Document(filePath)
        for para in doc_object.paragraphs:
            self.paragraphs.append(para.text)

    def get_whole_text(self):
        """
        Gets the file text and retuns it as a
        single string, with \n for paragraphs.
        """
        paras = self.get_paragraphs()
        wholeText = '\n'.join(paras)
