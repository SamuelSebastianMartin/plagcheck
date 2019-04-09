#!/usr/bin/env python3

#              reading_pack.txt and essay.txt

"""
This finds text in an essay.docx, which has bee copied from reading-pack.docx.
You will be asked to select, first a reading-pack document, then the essay.
The essay is processed by paragraph for matches.
Output is a new MS Word document with copied text in bold type.
"""

import docx
import os

import get_texts
import find_matches
import prune_matches
from span_class import Span


def main():
    # Get the texts
    rp_doc, rp, rp_path = get_texts.get_texts()
    sa_doc, sa, sa_path = get_texts.get_texts()

    #  Create final out document.
    out_doc = docx.Document()

    #  Loop through paragraphs, finding index of plagiarised text.
    for para in sa_doc.paragraphs:
        text = para.text
        para_words = get_texts.text_to_wordlist(text)

        # Find all plagiarised text, and return list of spans in 'para_words'
        matches = find_matches.find_matches(rp, para_words)
        print('Matches: ', matches)###

        # Order and sort spans to avoid overlaps and nested spans etc.
        indices = prune_matches.prune_indices(matches)
        print('after pruning :', indices)###
        span_objs = []

        # Turn 'para_words' list-spans into corresponding 'text' string-spans.
        sa_spans = []
        for span in indices:
            sp = Span(sa, para_words, span)
            sa_span = (sp.i, sp.j)
            sa_spans.append(sa_span)
        print(sa_spans)###

        out_doc = write_para(out_doc, text, sa_spans)

    out_doc.save('OUT.docx')
    os.system('libreoffice OUT.docx')


def find_in_sa(sa, matches, para_words):  ## Change to work with spans, not regex.
    """
    Takes the list-spans that tested positive in the reading pack, and finds
    the corresponding text-spans in the essay text.
    These text-spans are returned as a list.
    """
    spans = []
    n = 0  # Start point for search.
    for match in matches:
#        srch = match.search(sa, pos=n)
#        if srch != None:
#            span = srch.span()
#            n = span[0] +1  ## can change to end of last span b.. already sorted?
            spans.append(span)

    return spans

def write_para(out_doc, text, indices):
    p = out_doc.add_paragraph()

    i = 0
    j = 0
    for span in indices:
        j = span[0]  # End of unplagiarised section.
        p.add_run(text[i: j])
        p.add_run(text[span[0]: span[1]]).underline = True
        if span[1] < len(text):
            i = span[1]

    p.add_run(text[i: len(text)])  # Tail end of good text.

    return out_doc


    pass

if __name__ == '__main__':
    main()
