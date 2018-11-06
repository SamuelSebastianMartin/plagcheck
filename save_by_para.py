#! /usr/bin/env python3

'''
This is a shell that opens a docx file, processes it by paragraph
and returns a new docx file with the same text as the original.
The process() function is blank, but in that function will exist
the plagcheck process, and some way of formatting the paragraph
in runs
'''

import docx
import os


def open_text(textname):
    in_doc = docx.Document(textname)  # Create a Document object.
    return in_doc


def process(para):
    return para


def main():
    out_doc = docx.Document()
    in_doc = open_text('essay_long.docx')
    for para in in_doc.paragraphs:
        process(para)
        out_doc.add_paragraph(para.text)
    out_doc.save('para.docx')
    os.system('libreoffice para.docx')


if __name__ == '__main__':
    main()
