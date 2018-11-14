#! /usr/bin/env/ python3

import unittest
import get_texts
import os
import docx


class TestAdd(unittest.TestCase):

    def setUp(self):
        # Make new .docx file to open for testing.
        doc = docx.Document()
        doc.add_paragraph('Suck my Git.')
        doc.add_paragraph('Suck my doc.')
        doc.save('deleteme_testfile.docx')

    def test_open_text(self):
        doc_object = get_texts.open_text('deleteme_testfile.docx')
        d_type = type(doc_object)

        doc = docx.Document()
        document_type = type(doc)

        self.assertEqual(d_type, document_type)
        # There are 2 paragraphs.
        L = len(doc_object.paragraphs)
        self.assertEqual(L, 2)

    def test_doctotext(self):
        doc_object = get_texts.open_text('deleteme_testfile.docx')
        text = get_texts.doctotext(doc_object)
        result = type(text)
        str_type = type('random string')
        self.assertEqual(result, str_type)

    def test_prepare_text(self):
        text = 'The. fiSH & chips!\n me@mail!!! (3): www.dog'
        out = ['the', 'fish', 'chips', 'memail', '3', 'wwwdog']
        result = get_texts.prepare_text(text)
        self.assertEqual(result, out)

    def tearDown(self):
        os.remove('deleteme_testfile.docx')  # Clean up directory.


if __name__ == '__main__':
    unittest.main()
